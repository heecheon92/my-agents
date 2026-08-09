"""Group knowledge publish request API tests."""

from __future__ import annotations

from io import BytesIO

from docx import Document as WordDocument
from fastapi.testclient import TestClient
from openpyxl import Workbook
from sqlalchemy import select

from my_agents.api import create_app
from my_agents.knowledge.models import (
    DocumentChunkModel,
    DocumentModel,
    DocumentParseArtifactModel,
    KnowledgeBaseModel,
    KnowledgeBasePublicationModel,
    KnowledgePublishRequestModel,
)
from my_agents.knowledge.retrieval import RetrievalService
from my_agents.persistence.database import get_database_session

from .conftest import latest_auth_email_token, verify_latest_auth_email


def _client(
    monkeypatch,  # noqa: ANN001 - pytest monkeypatch fixture
    *,
    raise_server_exceptions: bool = True,
) -> TestClient:
    monkeypatch.setenv("MY_AGENTS_RESPONSE_MODE", "deterministic")
    monkeypatch.setenv("MY_AGENTS_SESSION_COOKIE_SECURE", "false")
    return TestClient(create_app(), raise_server_exceptions=raise_server_exceptions)


def _signup_login(client: TestClient, email: str) -> str:
    password = "correct horse battery staple"
    signup = client.post(
        "/auth/signup", json={"email": email, "nickname": "Test User", "password": password}
    )
    assert signup.status_code == 201
    verify_latest_auth_email(client, email)
    login = client.post("/auth/login", json={"email": email, "password": password})
    assert login.status_code == 200
    return signup.json()["user"]["id"]


def _create_group(owner: TestClient, *, name: str = "Publish Group") -> str:
    response = owner.post("/groups", json={"name": name})
    assert response.status_code == 201
    return response.json()["id"]


def _invite_and_accept_member(
    *,
    owner: TestClient,
    recipient: TestClient,
    group_id: str,
    recipient_email: str,
    role: str = "viewer",
) -> None:
    invitation = owner.post(
        f"/groups/{group_id}/invitations",
        json={"email": recipient_email, "role": role},
    )
    assert invitation.status_code == 201
    token = latest_auth_email_token(recipient_email, "group_invitation")
    accepted = recipient.post("/group-invitations/accept", json={"token": token})
    assert accepted.status_code == 200


def _create_personal_kb(client: TestClient, name: str = "Personal KB") -> str:
    response = client.post("/knowledge-bases", json={"name": name, "scope": "personal"})
    assert response.status_code == 201
    return response.json()["id"]


def _create_group_kb(client: TestClient, group_id: str, name: str = "Group KB") -> str:
    response = client.post(
        "/knowledge-bases",
        json={"name": name, "scope": "group", "group_id": group_id},
    )
    assert response.status_code == 201
    return response.json()["id"]


def _create_document(client: TestClient, *, kb_id: str, title: str, content: str) -> str:
    response = client.post(
        f"/knowledge-bases/{kb_id}/documents",
        json={"title": title, "content": content},
    )
    assert response.status_code == 201
    return response.json()["id"]


def _upload_xlsx_document(client: TestClient, *, kb_id: str, title: str) -> str:
    response = client.post(
        f"/knowledge-bases/{kb_id}/documents/upload",
        data={"title": title},
        files={
            "file": (
                "publish-metrics.xlsx",
                _xlsx_bytes(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )
    assert response.status_code == 201
    return response.json()["id"]


def _upload_docx_document(client: TestClient, *, kb_id: str, title: str) -> str:
    response = client.post(
        f"/knowledge-bases/{kb_id}/documents/upload",
        data={"title": title},
        files={
            "file": (
                "publish-word.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201
    return response.json()["id"]


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Publish Metrics"
    sheet.append(["Metric", "Value"])
    sheet.append(["GKPublishOfficeArtifact", "present"])
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _docx_bytes() -> bytes:
    document = WordDocument()
    document.add_heading("Publish Word Plan", level=1)
    document.add_paragraph("GKPublishWordArtifact is present in DOCX.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _ingest(client: TestClient, *, kb_id: str, document_id: str) -> None:
    response = client.post(f"/knowledge-bases/{kb_id}/documents/{document_id}/ingest")
    assert response.status_code == 200
    assert response.json()["status"] == "completed"


def _publish_personal_document(
    *,
    owner: TestClient,
    requester: TestClient,
    group_id: str,
    target_kb_id: str,
    personal_kb_id: str,
    title: str,
    content: str,
) -> str:
    source_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title=title,
        content=content,
    )
    publish_request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={
            "source_document_id": source_document_id,
            "target_knowledge_base_id": target_kb_id,
        },
    )
    assert publish_request.status_code == 201
    approved = owner.post(
        f"/groups/{group_id}/publish-requests/{publish_request.json()['id']}/approve"
    )
    assert approved.status_code == 200
    published_document_id = approved.json()["published_document_id"]
    assert published_document_id
    return published_document_id


def _group_retrieval_hits(user_id: str, *, kb_id: str, query: str) -> list[str]:
    session_generator = get_database_session()
    db = next(session_generator)
    try:
        return [
            item.document.id
            for item in RetrievalService(db).retrieve_scoped(
                user_id=user_id,
                query=query,
                knowledge_base_ids=[kb_id],
            )
        ]
    finally:
        session_generator.close()


def _retrieval_hits(user_id: str, *, kb_ids: list[str], query: str) -> list[str]:
    session_generator = get_database_session()
    db = next(session_generator)
    try:
        return [
            item.document.id
            for item in RetrievalService(db).retrieve_scoped(
                user_id=user_id,
                query=query,
                knowledge_base_ids=kb_ids,
            )
        ]
    finally:
        session_generator.close()


def test_member_can_create_pending_publish_request_for_own_personal_document(monkeypatch) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    _owner_id = _signup_login(owner, "publish-owner@example.com")
    requester_id = _signup_login(requester, "publish-requester@example.com")
    group_id = _create_group(owner)
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-requester@example.com",
        role="viewer",
    )
    target_kb_id = _create_group_kb(owner, group_id)
    personal_kb_id = _create_personal_kb(requester)
    source_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Private draft",
        content="Privately held publication candidate.",
    )

    response = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )

    assert response.status_code == 201
    payload = response.json()
    assert payload["requester_user_id"] == requester_id
    assert payload["target_group_id"] == group_id
    assert payload["target_knowledge_base_id"] == target_kb_id
    assert payload["source_document_id"] == source_document_id
    assert payload["source_knowledge_base_id"] is None
    assert payload["source_document_title"] == "Private draft"
    assert payload["source_document_excerpt"] == "Privately held publication candidate."
    assert payload["source_document_filename"] is None
    assert payload["source_knowledge_base_name"] is None
    assert payload["target_knowledge_base_name"] == "Group KB"
    assert payload["status"] == "pending"
    assert payload["reviewer_user_id"] is None
    assert payload["published_document_id"] is None
    assert payload["published_knowledge_base_id"] is None

    requester_list = requester.get(f"/groups/{group_id}/publish-requests")
    assert requester_list.status_code == 200
    assert [item["id"] for item in requester_list.json()] == [payload["id"]]

    owner_list = owner.get(f"/groups/{group_id}/publish-requests")
    assert owner_list.status_code == 200
    owner_payload = owner_list.json()[0]
    assert owner_payload["id"] == payload["id"]
    assert owner_payload["source_document_title"] == "Private draft"
    assert owner_payload["source_document_excerpt"] == "Privately held publication candidate."
    assert owner_payload["target_knowledge_base_name"] == "Group KB"


def test_owner_can_view_pending_publish_request_source_content(monkeypatch) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    viewer = _client(monkeypatch)
    _signup_login(owner, "publish-source-owner@example.com")
    _signup_login(requester, "publish-source-requester@example.com")
    _signup_login(viewer, "publish-source-viewer@example.com")
    group_id = _create_group(owner, name="Source Review Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-source-requester@example.com",
        role="viewer",
    )
    _invite_and_accept_member(
        owner=owner,
        recipient=viewer,
        group_id=group_id,
        recipient_email="publish-source-viewer@example.com",
        role="viewer",
    )
    target_kb_id = _create_group_kb(owner, group_id, "Source Review KB")
    personal_kb_id = _create_personal_kb(requester, "Source Review Personal KB")
    source_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Review me",
        content="Full extracted content that the owner must inspect before approval.",
    )
    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )
    assert request.status_code == 201
    request_id = request.json()["id"]

    viewer_source = viewer.get(f"/groups/{group_id}/publish-requests/{request_id}/source")
    assert viewer_source.status_code == 403

    owner_source = owner.get(f"/groups/{group_id}/publish-requests/{request_id}/source")
    assert owner_source.status_code == 200
    payload = owner_source.json()
    assert payload["request_id"] == request_id
    assert payload["source_kind"] == "document"
    assert payload["source_knowledge_base_id"] is None
    assert len(payload["documents"]) == 1
    [document] = payload["documents"]
    assert document["id"] == source_document_id
    assert document["title"] == "Review me"
    assert (
        document["content"] == "Full extracted content that the owner must inspect before approval."
    )
    assert document["source_type"] == "text"

    rejected = owner.post(f"/groups/{group_id}/publish-requests/{request_id}/reject")
    assert rejected.status_code == 200
    reviewed_source = owner.get(f"/groups/{group_id}/publish-requests/{request_id}/source")
    assert reviewed_source.status_code == 409


def test_owner_can_view_pending_whole_kb_publish_request_documents(monkeypatch) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    _signup_login(owner, "publish-kb-source-owner@example.com")
    _signup_login(requester, "publish-kb-source-requester@example.com")
    group_id = _create_group(owner, name="Whole KB Source Review Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-kb-source-requester@example.com",
        role="viewer",
    )
    personal_kb_id = _create_personal_kb(requester, "Whole KB For Review")
    first_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="First source",
        content="First full source body.",
    )
    second_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Second source",
        content="Second full source body.",
    )

    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_knowledge_base_id": personal_kb_id},
    )
    assert request.status_code == 201
    request_id = request.json()["id"]

    owner_source = owner.get(f"/groups/{group_id}/publish-requests/{request_id}/source")
    assert owner_source.status_code == 200
    payload = owner_source.json()
    assert payload["source_kind"] == "knowledge_base"
    assert payload["source_knowledge_base_id"] == personal_kb_id
    assert payload["source_knowledge_base_name"] == "Whole KB For Review"
    assert [document["id"] for document in payload["documents"]] == [
        first_document_id,
        second_document_id,
    ]
    assert [document["title"] for document in payload["documents"]] == [
        "First source",
        "Second source",
    ]
    assert [document["content"] for document in payload["documents"]] == [
        "First full source body.",
        "Second full source body.",
    ]


def test_requester_can_cancel_pending_document_publish_request_and_request_again(
    monkeypatch,
) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    _signup_login(owner, "publish-cancel-doc-owner@example.com")
    _signup_login(requester, "publish-cancel-doc-requester@example.com")
    group_id = _create_group(owner, name="Cancel Document Request Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-cancel-doc-requester@example.com",
        role="viewer",
    )
    target_kb_id = _create_group_kb(owner, group_id, "Cancel Document Group KB")
    personal_kb_id = _create_personal_kb(requester, "Cancel Document Personal KB")
    source_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Cancelable draft",
        content="Requester can cancel before review.",
    )
    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )
    assert request.status_code == 201
    request_id = request.json()["id"]

    owner_cancel = owner.post(f"/groups/{group_id}/publish-requests/{request_id}/cancel")
    assert owner_cancel.status_code == 403

    cancelled = requester.post(f"/groups/{group_id}/publish-requests/{request_id}/cancel")

    assert cancelled.status_code == 200
    assert cancelled.json()["status"] == "cancelled"
    assert cancelled.json()["source_document_id"] == source_document_id
    second_request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )
    assert second_request.status_code == 201
    assert second_request.json()["id"] != request_id
    approve_cancelled = owner.post(f"/groups/{group_id}/publish-requests/{request_id}/approve")
    assert approve_cancelled.status_code == 409
    assert approve_cancelled.json()["code"] == "publish_request_already_reviewed"


def test_requester_can_cancel_pending_whole_kb_publish_request_and_request_again(
    monkeypatch,
) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    _signup_login(owner, "publish-cancel-kb-owner@example.com")
    _signup_login(requester, "publish-cancel-kb-requester@example.com")
    group_id = _create_group(owner, name="Cancel Whole KB Request Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-cancel-kb-requester@example.com",
        role="viewer",
    )
    personal_kb_id = _create_personal_kb(requester, "Cancelable Whole KB")
    _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Cancelable KB source",
        content="Requester can cancel whole KB share before review.",
    )
    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_knowledge_base_id": personal_kb_id},
    )
    assert request.status_code == 201
    request_id = request.json()["id"]
    duplicate_pending = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_knowledge_base_id": personal_kb_id},
    )
    assert duplicate_pending.status_code == 409

    cancelled = requester.post(f"/groups/{group_id}/publish-requests/{request_id}/cancel")

    assert cancelled.status_code == 200
    assert cancelled.json()["status"] == "cancelled"
    second_request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_knowledge_base_id": personal_kb_id},
    )
    assert second_request.status_code == 201
    assert second_request.json()["id"] != request_id


def test_publish_request_rejects_invalid_source_or_target_boundaries(monkeypatch) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    other = _client(monkeypatch)
    _owner_id = _signup_login(owner, "publish-boundary-owner@example.com")
    _requester_id = _signup_login(requester, "publish-boundary-requester@example.com")
    _other_id = _signup_login(other, "publish-boundary-other@example.com")
    group_id = _create_group(owner, name="Boundary Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-boundary-requester@example.com",
        role="viewer",
    )
    target_group_kb_id = _create_group_kb(owner, group_id, "Boundary Group KB")
    requester_personal_kb_id = _create_personal_kb(requester, "Requester KB")
    other_personal_kb_id = _create_personal_kb(other, "Other KB")
    requester_personal_doc_id = _create_document(
        requester,
        kb_id=requester_personal_kb_id,
        title="Requester personal",
        content="requester owned",
    )
    other_personal_doc_id = _create_document(
        other,
        kb_id=other_personal_kb_id,
        title="Other personal",
        content="other owned",
    )
    owner_personal_kb_id = _create_personal_kb(owner, "Owner Source KB")
    group_doc_id = _publish_personal_document(
        owner=owner,
        requester=owner,
        group_id=group_id,
        target_kb_id=target_group_kb_id,
        personal_kb_id=owner_personal_kb_id,
        title="Group source",
        content="already group scoped",
    )
    requester_non_group_target_kb_id = _create_personal_kb(requester, "Not group target")

    not_owned = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={
            "source_document_id": other_personal_doc_id,
            "target_knowledge_base_id": target_group_kb_id,
        },
    )
    assert not_owned.status_code == 404

    non_personal_source = owner.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": group_doc_id, "target_knowledge_base_id": target_group_kb_id},
    )
    assert non_personal_source.status_code == 422
    assert non_personal_source.json()["detail"] == "source document must be personal"

    unsupported_hidden_kb_id = _create_personal_kb(requester, "Unsupported Hidden KB")
    unsupported_hidden_doc_id = _create_document(
        requester,
        kb_id=unsupported_hidden_kb_id,
        title="Unsupported hidden",
        content="unsupported hidden source should not publish by raw id",
    )
    session_generator = get_database_session()
    db = next(session_generator)
    try:
        unsupported_hidden_kb = db.get(KnowledgeBaseModel, unsupported_hidden_kb_id)
        assert unsupported_hidden_kb is not None
        unsupported_hidden_kb.purpose = "unsupported_hidden"
        db.add(unsupported_hidden_kb)
        db.commit()
    finally:
        session_generator.close()
    hidden_nonstandard_source = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={
            "source_document_id": unsupported_hidden_doc_id,
            "target_knowledge_base_id": target_group_kb_id,
        },
    )
    assert hidden_nonstandard_source.status_code == 422
    assert hidden_nonstandard_source.json()["detail"] == "source document must be personal"

    non_group_target = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={
            "source_document_id": requester_personal_doc_id,
            "target_knowledge_base_id": requester_non_group_target_kb_id,
        },
    )
    assert non_group_target.status_code == 422
    assert non_group_target.json()["detail"] == "target knowledge base must belong to group"

    both_sources = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={
            "source_document_id": requester_personal_doc_id,
            "source_knowledge_base_id": requester_personal_kb_id,
            "target_knowledge_base_id": target_group_kb_id,
        },
    )
    assert both_sources.status_code == 422

    missing_target = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": requester_personal_doc_id},
    )
    assert missing_target.status_code == 422

    group_kb_source = owner.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_knowledge_base_id": target_group_kb_id},
    )
    assert group_kb_source.status_code == 422
    assert group_kb_source.json()["detail"] == "source knowledge base must be personal"


def test_owner_admin_approval_copies_and_ingests_group_document_without_exposing_original(
    monkeypatch,
) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    viewer = _client(monkeypatch)
    owner_id = _signup_login(owner, "publish-approve-owner@example.com")
    _requester_id = _signup_login(requester, "publish-approve-requester@example.com")
    viewer_id = _signup_login(viewer, "publish-approve-viewer@example.com")
    group_id = _create_group(owner, name="Approval Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-approve-requester@example.com",
        role="viewer",
    )
    _invite_and_accept_member(
        owner=owner,
        recipient=viewer,
        group_id=group_id,
        recipient_email="publish-approve-viewer@example.com",
        role="viewer",
    )
    target_kb_id = _create_group_kb(owner, group_id, "Approval KB")
    personal_kb_id = _create_personal_kb(requester, "Approval Personal KB")
    source_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Personal Alpha",
        content="GKPublishUniqueAlpha is requester-only until approval.",
    )
    _ingest(requester, kb_id=personal_kb_id, document_id=source_document_id)

    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )
    assert request.status_code == 201
    request_id = request.json()["id"]

    assert _group_retrieval_hits(viewer_id, kb_id=target_kb_id, query="GKPublishUniqueAlpha") == []

    viewer_approve = viewer.post(f"/groups/{group_id}/publish-requests/{request_id}/approve")
    assert viewer_approve.status_code == 403
    assert _group_retrieval_hits(viewer_id, kb_id=target_kb_id, query="GKPublishUniqueAlpha") == []

    approved = owner.post(f"/groups/{group_id}/publish-requests/{request_id}/approve")
    assert approved.status_code == 200
    approved_payload = approved.json()
    assert approved_payload["status"] == "approved"
    assert approved_payload["reviewer_user_id"] == owner_id
    published_document_id = approved_payload["published_document_id"]
    assert published_document_id and published_document_id != source_document_id

    hits = _group_retrieval_hits(viewer_id, kb_id=target_kb_id, query="GKPublishUniqueAlpha")
    assert hits == [published_document_id]

    requester_delete_copy = requester.delete(f"/documents/{published_document_id}")
    assert requester_delete_copy.status_code == 404
    requester_manage_copy = requester.patch(
        f"/documents/{published_document_id}/permissions",
        json={
            "user_id": viewer_id,
            "can_read": True,
            "can_write": True,
            "can_manage": True,
            "can_ingest": True,
        },
    )
    assert requester_manage_copy.status_code == 403
    owner_manage_copy = owner.patch(
        f"/documents/{published_document_id}/permissions",
        json={
            "user_id": viewer_id,
            "can_read": True,
            "can_write": False,
            "can_manage": False,
            "can_ingest": False,
        },
    )
    assert owner_manage_copy.status_code == 200

    session_generator = get_database_session()
    db = next(session_generator)
    try:
        original = db.get(DocumentModel, source_document_id)
        copied = db.get(DocumentModel, published_document_id)
        publish_request = db.scalar(
            select(KnowledgePublishRequestModel).where(
                KnowledgePublishRequestModel.id == request_id
            )
        )
        assert original is not None
        assert copied is not None
        assert publish_request is not None
        assert original.group_id is None
        assert original.knowledge_base_id == personal_kb_id
        assert copied.group_id == group_id
        assert copied.knowledge_base_id == target_kb_id
        assert copied.content == original.content
        assert publish_request.published_document_id == published_document_id
    finally:
        session_generator.close()


def test_owner_can_delete_published_group_copy_without_breaking_request_history(
    monkeypatch,
) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    owner_id = _signup_login(owner, "publish-delete-copy-owner@example.com")
    _requester_id = _signup_login(requester, "publish-delete-copy-requester@example.com")
    group_id = _create_group(owner, name="Publish Delete Copy Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-delete-copy-requester@example.com",
        role="viewer",
    )
    target_kb_id = _create_group_kb(owner, group_id, "Delete Copy Group KB")
    personal_kb_id = _create_personal_kb(requester, "Delete Copy Personal KB")
    source_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Delete Published Copy Source",
        content="GKDeletePublishedCopyAlpha should disappear after copy deletion.",
    )
    _ingest(requester, kb_id=personal_kb_id, document_id=source_document_id)
    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )
    assert request.status_code == 201
    request_id = request.json()["id"]
    approved = owner.post(f"/groups/{group_id}/publish-requests/{request_id}/approve")
    assert approved.status_code == 200
    published_document_id = approved.json()["published_document_id"]
    assert published_document_id
    assert _group_retrieval_hits(
        owner_id,
        kb_id=target_kb_id,
        query="GKDeletePublishedCopyAlpha",
    ) == [published_document_id]

    response = owner.delete(f"/documents/{published_document_id}")

    assert response.status_code == 204
    assert (
        _group_retrieval_hits(
            owner_id,
            kb_id=target_kb_id,
            query="GKDeletePublishedCopyAlpha",
        )
        == []
    )
    listed = owner.get(f"/groups/{group_id}/publish-requests")
    assert listed.status_code == 200
    payload = listed.json()[0]
    assert payload["id"] == request_id
    assert payload["status"] == "approved"
    assert payload["source_document_id"] == source_document_id
    assert payload["source_document_title"] == "Delete Published Copy Source"
    assert payload["published_document_id"] is None
    session_generator = get_database_session()
    db = next(session_generator)
    try:
        publish_request = db.scalar(
            select(KnowledgePublishRequestModel).where(
                KnowledgePublishRequestModel.id == request_id
            )
        )
        assert db.get(DocumentModel, published_document_id) is None
        assert publish_request is not None
        assert publish_request.source_document_id == source_document_id
        assert publish_request.published_document_id is None
    finally:
        session_generator.close()


def test_pending_publish_request_is_withdrawn_when_source_document_is_deleted(
    monkeypatch,
) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    _owner_id = _signup_login(owner, "publish-delete-source-owner@example.com")
    _requester_id = _signup_login(requester, "publish-delete-source-requester@example.com")
    group_id = _create_group(owner, name="Publish Delete Source Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-delete-source-requester@example.com",
        role="viewer",
    )
    target_kb_id = _create_group_kb(owner, group_id, "Delete Source Group KB")
    personal_kb_id = _create_personal_kb(requester, "Delete Source Personal KB")
    source_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Delete Source Candidate",
        content="Deleting a source should withdraw pending publish review.",
    )
    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )
    assert request.status_code == 201
    request_id = request.json()["id"]

    response = requester.delete(f"/documents/{source_document_id}")

    assert response.status_code == 204
    listed = owner.get(f"/groups/{group_id}/publish-requests")
    assert listed.status_code == 200
    payload = listed.json()[0]
    assert payload["id"] == request_id
    assert payload["status"] == "withdrawn"
    assert payload["source_document_id"] is None
    assert payload["source_document_title"] == "Delete Source Candidate"
    assert payload["source_document_excerpt"] == (
        "Deleting a source should withdraw pending publish review."
    )
    assert payload["published_document_id"] is None
    approve_after_delete = owner.post(f"/groups/{group_id}/publish-requests/{request_id}/approve")
    assert approve_after_delete.status_code == 409
    session_generator = get_database_session()
    db = next(session_generator)
    try:
        publish_request = db.scalar(
            select(KnowledgePublishRequestModel).where(
                KnowledgePublishRequestModel.id == request_id
            )
        )
        assert db.get(DocumentModel, source_document_id) is None
        assert publish_request is not None
        assert publish_request.status == "withdrawn"
        assert publish_request.source_document_id is None
        assert publish_request.published_document_id is None
    finally:
        session_generator.close()


def test_owner_approval_copies_office_parse_artifact_to_group_document(monkeypatch) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    _owner_id = _signup_login(owner, "publish-office-owner@example.com")
    _requester_id = _signup_login(requester, "publish-office-requester@example.com")
    group_id = _create_group(owner, name="Office Publish Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-office-requester@example.com",
        role="viewer",
    )
    target_kb_id = _create_group_kb(owner, group_id, "Office Publish KB")
    personal_kb_id = _create_personal_kb(requester, "Office Publish Personal KB")
    source_document_id = _upload_xlsx_document(
        requester,
        kb_id=personal_kb_id,
        title="Publish Metrics Workbook",
    )

    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )
    assert request.status_code == 201

    approved = owner.post(f"/groups/{group_id}/publish-requests/{request.json()['id']}/approve")

    assert approved.status_code == 200
    published_document_id = approved.json()["published_document_id"]
    assert published_document_id and published_document_id != source_document_id

    session_generator = get_database_session()
    db = next(session_generator)
    try:
        source_artifact = db.scalar(
            select(DocumentParseArtifactModel).where(
                DocumentParseArtifactModel.document_id == source_document_id
            )
        )
        copied_artifact = db.scalar(
            select(DocumentParseArtifactModel).where(
                DocumentParseArtifactModel.document_id == published_document_id
            )
        )
        assert source_artifact is not None
        assert copied_artifact is not None
        assert copied_artifact.id != source_artifact.id
        assert copied_artifact.markdown_content == source_artifact.markdown_content
        assert copied_artifact.elements_json == source_artifact.elements_json
        assert copied_artifact.source_filename == "publish-metrics.xlsx"
        assert copied_artifact.source_sha256 == source_artifact.source_sha256
    finally:
        session_generator.close()


def test_owner_approval_copies_docx_parse_artifact_to_group_document(monkeypatch) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    _owner_id = _signup_login(owner, "publish-docx-owner@example.com")
    _requester_id = _signup_login(requester, "publish-docx-requester@example.com")
    group_id = _create_group(owner, name="DOCX Publish Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-docx-requester@example.com",
        role="viewer",
    )
    target_kb_id = _create_group_kb(owner, group_id, "DOCX Publish KB")
    personal_kb_id = _create_personal_kb(requester, "DOCX Publish Personal KB")
    source_document_id = _upload_docx_document(
        requester,
        kb_id=personal_kb_id,
        title="Publish Word Document",
    )

    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )
    assert request.status_code == 201

    approved = owner.post(f"/groups/{group_id}/publish-requests/{request.json()['id']}/approve")

    assert approved.status_code == 200
    published_document_id = approved.json()["published_document_id"]
    assert published_document_id and published_document_id != source_document_id

    session_generator = get_database_session()
    db = next(session_generator)
    try:
        source_artifact = db.scalar(
            select(DocumentParseArtifactModel).where(
                DocumentParseArtifactModel.document_id == source_document_id
            )
        )
        copied_artifact = db.scalar(
            select(DocumentParseArtifactModel).where(
                DocumentParseArtifactModel.document_id == published_document_id
            )
        )
        assert source_artifact is not None
        assert copied_artifact is not None
        assert copied_artifact.id != source_artifact.id
        assert copied_artifact.markdown_content == source_artifact.markdown_content
        assert copied_artifact.elements_json == source_artifact.elements_json
        assert copied_artifact.source_filename == "publish-word.docx"
        assert copied_artifact.source_type == "word_document"
        assert copied_artifact.source_sha256 == source_artifact.source_sha256
    finally:
        session_generator.close()


def test_publish_approval_rolls_back_group_copy_when_ingestion_fails(monkeypatch) -> None:  # noqa: ANN001
    def fail_ingestion(self, document):  # noqa: ANN001, ARG001
        raise RuntimeError("fixture ingestion failure")

    monkeypatch.setattr(
        "my_agents.api.groups.KnowledgeExtractionService.ingest_document",
        fail_ingestion,
    )
    owner = _client(monkeypatch, raise_server_exceptions=False)
    requester = _client(monkeypatch, raise_server_exceptions=False)
    _owner_id = _signup_login(owner, "publish-ingest-fail-owner@example.com")
    _requester_id = _signup_login(requester, "publish-ingest-fail-requester@example.com")
    group_id = _create_group(owner, name="Ingest Failure Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-ingest-fail-requester@example.com",
        role="viewer",
    )
    target_kb_id = _create_group_kb(owner, group_id, "Ingest Failure KB")
    personal_kb_id = _create_personal_kb(requester, "Ingest Failure Personal KB")
    source_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Rollback Source",
        content="Publish rollback should not leave an approved group copy.",
    )
    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )
    assert request.status_code == 201
    request_id = request.json()["id"]

    approved = owner.post(f"/groups/{group_id}/publish-requests/{request_id}/approve")

    assert approved.status_code == 500
    session_generator = get_database_session()
    db = next(session_generator)
    try:
        publish_request = db.scalar(
            select(KnowledgePublishRequestModel).where(
                KnowledgePublishRequestModel.id == request_id
            )
        )
        leaked_group_copies = db.scalars(
            select(DocumentModel).where(
                DocumentModel.knowledge_base_id == target_kb_id,
                DocumentModel.title == "Rollback Source",
            )
        ).all()
        assert publish_request is not None
        assert publish_request.status == "pending"
        assert publish_request.published_document_id is None
        assert leaked_group_copies == []
    finally:
        session_generator.close()


def test_rejected_publish_request_has_zero_retrieval_effect(monkeypatch) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    viewer = _client(monkeypatch)
    owner_id = _signup_login(owner, "publish-reject-owner@example.com")
    _requester_id = _signup_login(requester, "publish-reject-requester@example.com")
    viewer_id = _signup_login(viewer, "publish-reject-viewer@example.com")
    group_id = _create_group(owner, name="Reject Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-reject-requester@example.com",
        role="viewer",
    )
    _invite_and_accept_member(
        owner=owner,
        recipient=viewer,
        group_id=group_id,
        recipient_email="publish-reject-viewer@example.com",
        role="viewer",
    )
    target_kb_id = _create_group_kb(owner, group_id, "Reject KB")
    personal_kb_id = _create_personal_kb(requester, "Reject Personal KB")
    source_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Personal Beta",
        content="GKPublishUniqueBeta should never appear in group retrieval.",
    )
    _ingest(requester, kb_id=personal_kb_id, document_id=source_document_id)

    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_document_id": source_document_id, "target_knowledge_base_id": target_kb_id},
    )
    assert request.status_code == 201
    request_id = request.json()["id"]

    rejected = owner.post(f"/groups/{group_id}/publish-requests/{request_id}/reject")
    assert rejected.status_code == 200
    assert rejected.json()["status"] == "rejected"
    assert rejected.json()["reviewer_user_id"] == owner_id
    assert rejected.json()["published_document_id"] is None

    assert _group_retrieval_hits(viewer_id, kb_id=target_kb_id, query="GKPublishUniqueBeta") == []


def test_owner_approval_publishes_entire_personal_kb_as_group_source(monkeypatch) -> None:  # noqa: ANN001
    owner = _client(monkeypatch)
    requester = _client(monkeypatch)
    viewer = _client(monkeypatch)
    owner_id = _signup_login(owner, "publish-kb-owner@example.com")
    _requester_id = _signup_login(requester, "publish-kb-requester@example.com")
    viewer_id = _signup_login(viewer, "publish-kb-viewer@example.com")
    group_id = _create_group(owner, name="Published Personal KB Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-kb-requester@example.com",
        role="viewer",
    )
    _invite_and_accept_member(
        owner=owner,
        recipient=viewer,
        group_id=group_id,
        recipient_email="publish-kb-viewer@example.com",
        role="viewer",
    )
    personal_kb_id = _create_personal_kb(requester, "Requester Public Candidate KB")
    source_document_id = _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Whole Personal KB Source",
        content="GKPublishedPersonalKbAlpha should become a copied group source.",
    )
    source_docx_id = _upload_docx_document(
        requester,
        kb_id=personal_kb_id,
        title="Whole Personal KB Office Source",
    )
    _ingest(requester, kb_id=personal_kb_id, document_id=source_document_id)
    _ingest(requester, kb_id=personal_kb_id, document_id=source_docx_id)

    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_knowledge_base_id": personal_kb_id},
    )
    assert request.status_code == 201
    request_payload = request.json()
    assert request_payload["source_document_id"] is None
    assert request_payload["target_knowledge_base_id"] is None
    assert request_payload["source_knowledge_base_id"] == personal_kb_id
    assert request_payload["source_knowledge_base_name"] == "Requester Public Candidate KB"
    assert request_payload["source_document_title"] is None
    assert request_payload["target_knowledge_base_name"] is None
    assert request_payload["published_knowledge_base_id"] is None
    assert request_payload["published_knowledge_base_name"] is None
    assert (
        _retrieval_hits(
            viewer_id,
            kb_ids=[personal_kb_id],
            query="GKPublishedPersonalKbAlpha",
        )
        == []
    )

    viewer_approve = viewer.post(
        f"/groups/{group_id}/publish-requests/{request_payload['id']}/approve"
    )
    assert viewer_approve.status_code == 403

    approved = owner.post(f"/groups/{group_id}/publish-requests/{request_payload['id']}/approve")
    assert approved.status_code == 200
    approved_payload = approved.json()
    group_copy_id = approved_payload["published_knowledge_base_id"]
    assert approved_payload["status"] == "approved"
    assert approved_payload["reviewer_user_id"] == owner_id
    assert approved_payload["published_document_id"] is None
    assert group_copy_id and group_copy_id != personal_kb_id
    assert approved_payload["published_knowledge_base_name"] == "Requester Public Candidate KB"
    assert (
        _retrieval_hits(
            viewer_id,
            kb_ids=[personal_kb_id],
            query="GKPublishedPersonalKbAlpha",
        )
        == []
    )

    session_generator = get_database_session()
    db = next(session_generator)
    try:
        group_copy = db.get(KnowledgeBaseModel, group_copy_id)
        assert group_copy is not None
        assert group_copy.scope == "group"
        assert group_copy.group_id == group_id
        assert group_copy.purpose == "standard"
        copied_documents = db.scalars(
            select(DocumentModel).where(DocumentModel.knowledge_base_id == group_copy_id)
        ).all()
        copied_document_ids = {document.id for document in copied_documents}
        assert {document.title for document in copied_documents} == {
            "Whole Personal KB Source",
            "Whole Personal KB Office Source",
        }
        assert source_document_id not in copied_document_ids
        assert source_docx_id not in copied_document_ids
        copied_text_document = next(
            document
            for document in copied_documents
            if document.title == "Whole Personal KB Source"
        )
        copied_docx_document = next(
            document
            for document in copied_documents
            if document.title == "Whole Personal KB Office Source"
        )
        source_artifact = db.scalar(
            select(DocumentParseArtifactModel).where(
                DocumentParseArtifactModel.document_id == source_docx_id
            )
        )
        copied_artifact = db.scalar(
            select(DocumentParseArtifactModel).where(
                DocumentParseArtifactModel.document_id == copied_docx_document.id
            )
        )
        assert source_artifact is not None
        assert copied_artifact is not None
        assert copied_artifact.id != source_artifact.id
        assert copied_artifact.markdown_content == source_artifact.markdown_content
        assert copied_artifact.source_filename == "publish-word.docx"
        copied_chunks = db.scalars(
            select(DocumentChunkModel).where(
                DocumentChunkModel.document_id.in_(copied_document_ids)
            )
        ).all()
        assert copied_chunks
        assert {chunk.document_id for chunk in copied_chunks}.issubset(copied_document_ids)
        publication = db.scalar(
            select(KnowledgeBasePublicationModel).where(
                KnowledgeBasePublicationModel.group_id == group_id,
                KnowledgeBasePublicationModel.knowledge_base_id == personal_kb_id,
            )
        )
        assert publication is None
    finally:
        session_generator.close()

    assert _retrieval_hits(
        viewer_id,
        kb_ids=[group_copy_id],
        query="GKPublishedPersonalKbAlpha",
    ) == [copied_text_document.id]
    assert _retrieval_hits(
        viewer_id,
        kb_ids=[group_copy_id],
        query="GKPublishWordArtifact",
    ) == [copied_docx_document.id]

    requester_knowledge_bases = requester.get("/knowledge-bases").json()
    original_listing = next(
        knowledge_base
        for knowledge_base in requester_knowledge_bases
        if knowledge_base["id"] == personal_kb_id
    )
    group_copy_listing = next(
        knowledge_base
        for knowledge_base in requester_knowledge_bases
        if knowledge_base["id"] == group_copy_id
    )
    assert original_listing["published_group_ids"] == []
    assert group_copy_listing["group_id"] == group_id
    assert group_copy_listing["published_group_ids"] == [group_id]

    duplicate = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_knowledge_base_id": personal_kb_id},
    )
    assert duplicate.status_code == 409
    assert duplicate.json()["detail"] == "knowledge base already published to group"

    deleted_original = requester.delete(f"/knowledge-bases/{personal_kb_id}")
    assert deleted_original.status_code == 204
    assert _retrieval_hits(
        viewer_id,
        kb_ids=[group_copy_id],
        query="GKPublishedPersonalKbAlpha",
    ) == [copied_text_document.id]
    history_after_source_delete = owner.get(f"/groups/{group_id}/publish-requests").json()
    approved_history = next(
        item for item in history_after_source_delete if item["id"] == request_payload["id"]
    )
    assert approved_history["source_knowledge_base_id"] is None
    assert approved_history["source_knowledge_base_name"] == "Requester Public Candidate KB"
    assert approved_history["published_knowledge_base_id"] == group_copy_id
    assert approved_history["published_knowledge_base_name"] == "Requester Public Candidate KB"

    deleted_group_copy = owner.delete(f"/knowledge-bases/{group_copy_id}")
    assert deleted_group_copy.status_code == 204
    assert (
        _retrieval_hits(
            viewer_id,
            kb_ids=[group_copy_id],
            query="GKPublishedPersonalKbAlpha",
        )
        == []
    )
    history_after_group_delete = owner.get(f"/groups/{group_id}/publish-requests").json()
    deleted_copy_history = next(
        item for item in history_after_group_delete if item["id"] == request_payload["id"]
    )
    assert deleted_copy_history["status"] == "approved"
    assert deleted_copy_history["published_knowledge_base_id"] is None
    assert deleted_copy_history["published_knowledge_base_name"] == "Requester Public Candidate KB"


def test_whole_kb_publish_approval_rolls_back_group_copy_when_ingestion_fails(monkeypatch) -> None:  # noqa: ANN001
    def fail_ingestion(self, document):  # noqa: ANN001, ARG001
        raise RuntimeError("fixture whole-kb ingestion failure")

    monkeypatch.setattr(
        "my_agents.knowledge.publication_copies.KnowledgeExtractionService.ingest_document",
        fail_ingestion,
    )
    owner = _client(monkeypatch, raise_server_exceptions=False)
    requester = _client(monkeypatch, raise_server_exceptions=False)
    _owner_id = _signup_login(owner, "publish-kb-fail-owner@example.com")
    _requester_id = _signup_login(requester, "publish-kb-fail-requester@example.com")
    group_id = _create_group(owner, name="Whole KB Failure Group")
    _invite_and_accept_member(
        owner=owner,
        recipient=requester,
        group_id=group_id,
        recipient_email="publish-kb-fail-requester@example.com",
        role="viewer",
    )
    personal_kb_id = _create_personal_kb(requester, "Whole KB Failure Personal KB")
    _create_document(
        requester,
        kb_id=personal_kb_id,
        title="Whole KB Rollback Source",
        content="GKWholeKbRollback should not leak through a partial group copy.",
    )
    request = requester.post(
        f"/groups/{group_id}/publish-requests",
        json={"source_knowledge_base_id": personal_kb_id},
    )
    assert request.status_code == 201
    request_id = request.json()["id"]

    approved = owner.post(f"/groups/{group_id}/publish-requests/{request_id}/approve")

    assert approved.status_code == 500
    session_generator = get_database_session()
    db = next(session_generator)
    try:
        publish_request = db.scalar(
            select(KnowledgePublishRequestModel).where(
                KnowledgePublishRequestModel.id == request_id
            )
        )
        leaked_group_kbs = db.scalars(
            select(KnowledgeBaseModel).where(
                KnowledgeBaseModel.group_id == group_id,
                KnowledgeBaseModel.name.like("Whole KB Failure Personal KB%"),
            )
        ).all()
        leaked_documents = db.scalars(
            select(DocumentModel).where(DocumentModel.title == "Whole KB Rollback Source")
        ).all()
        assert publish_request is not None
        assert publish_request.status == "pending"
        assert publish_request.published_knowledge_base_id is None
        assert leaked_group_kbs == []
        assert all(document.knowledge_base_id == personal_kb_id for document in leaked_documents)
    finally:
        session_generator.close()
